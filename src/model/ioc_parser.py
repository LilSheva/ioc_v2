import re
from urllib.parse import urlparse
from docx import Document

class IOCParser:
    def _get_text_from_docx(self, filepath):
        try:
            document = Document(filepath)
            full_text = []
            for para in document.paragraphs: full_text.append(para.text)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells: full_text.append(cell.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Ошибка при чтении файла {filepath}: {e}"); return None

    def extract_all_iocs(self, filepaths, ioc_config):
        """
        Главный метод. Принимает СПИСОК путей к файлам и конфигурацию.
        """
        all_texts = []
        for path in filepaths:
            text = self._get_text_from_docx(path)
            if text:
                all_texts.append(text)
        
        if not all_texts:
            return {}

        content = '\n\n--- END OF FILE ---\n\n'.join(all_texts)
        
        raw_finds = {}
        final_data = {}
        all_enabled_types = {name: data for name, data in ioc_config.items() if data.get("enabled")}

        for name, data in all_enabled_types.items():
            try:
                regex = data.get("regex", "")
                if not regex: continue
                if '(?P<value>' in regex:
                    found = [match.group('value') for match in re.finditer(regex, content)]
                else:
                    found = re.findall(regex, content)
                if found: raw_finds[name] = list(set(found))
            except re.error as e: print(f"Ошибка в регулярном выражении для {name}: {e}")
        
        uris = raw_finds.get('URI', [])
        sha1s = raw_finds.get('SHA1', [])
        sha256s = raw_finds.get('SHA256', [])
        
        if 'DNS' in raw_finds:
            uri_domains = {urlparse(uri.replace('[.]', '.')).netloc for uri in uris}
            potential_dns = raw_finds['DNS']
            clean_dns = [
                item for item in potential_dns 
                if item.replace('[.]', '.') not in uri_domains 
                and not all(part.isdigit() for part in item.replace('[.]', '.').split('.'))
            ]
            if clean_dns: final_data['DNS'] = clean_dns
        
        if 'MD5' in raw_finds:
            sha_set = set(sha1s) | set(sha256s)
            potential_md5 = raw_finds['MD5']
            clean_md5 = [h for h in potential_md5 if not any(h in s for s in sha_set)]
            if clean_md5: final_data['MD5'] = clean_md5

        for name, found_items in raw_finds.items():
            if name not in ['DNS', 'MD5']:
                final_data[name] = found_items
                
        return final_data